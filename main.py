import os
import pandas as pd
from docx import Document
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service
import time

info = pd.read_csv("student_data.csv")
for index in range(info.shape[0]):
    doc = Document('test.docx')
    section = doc.sections[0]
    header = section.header
    for paragraph in header.paragraphs:
        paragraph.clear()

    new_paragraph = header.paragraphs[0]
    name = info["Name "][index]
    roll_no = info["Roll_no"][index]
    header = name + " "*(75-len(name)) + "CSE B"+ " "*50 + str(roll_no)
    new_paragraph.add_run(header)

    doc.save("named_files"+"\\"+name+'.docx')


options  = webdriver.ChromeOptions()
options.add_argument('--profile-directory=default')
options.add_argument(r"user-data-dir=C:\Users\admin\AppData\Local\Google\Chrome\User Data")
service = Service(executable_path=r"D:\User Data\Desktop\nishant\PYTHON\selenium\chromedriver.exe")
driver = webdriver.Chrome(service=service,options=options)

for index in range(info.shape[0]): 
    contact_name = str(info["number"][index])
    docx_file_path = r'D:\User Data\Desktop\Profile Project\Connected Projects\Automate assignment sender\named_files\\'+"\\"+info["Name "][index]+".docx"
    driver.get(f'https://web.whatsapp.com/send?phone={contact_name}')
    while True:
        try:
            attachment_box = driver.find_element("xpath",'//div[@title="Attach"]')
            attachment_box.click()
            time.sleep(3)
            break
        except: pass

    time.sleep(2)
    doc_button = driver.find_element("xpath",'//input[@accept="*"]')
    doc_button.send_keys(docx_file_path)
                

    time.sleep(3)
    send_button = driver.find_element("xpath",'//span[@data-icon="send"]')
    send_button.click()
    
    
driver.quit()
for path in os.listdir(r"named_files"):
    os.remove("named_files"+"\\"+path)




